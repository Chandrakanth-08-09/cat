import base64
import uuid
import os
import io
import datetime
from django.shortcuts import render, redirect
from django.conf import settings
from django.contrib import messages
from .forms import ImageUploadForm
from .models import ImageUpload
from docx import Document
from django.http import HttpResponse
from django.core.files.base import ContentFile
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError
from PIL import Image as PILImage
import numpy as np
import cv2
import google.generativeai as genai
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import pandas as pd
from docx2pdf import convert

GOOGLE_API_KEY = ''  # Replace with your actual API key

genai.configure(api_key=GOOGLE_API_KEY)
model = genai.GenerativeModel(model_name="gemini-1.5-pro-latest")

def upload_image(request):
    if request.method == 'POST':
        # Get the base64-encoded image data from the form
        image_data_url = request.POST.get('image_data')

        if image_data_url:
            try:
                # Decode base64 data
                format, imgstr = image_data_url.split(';base64,')
                ext = format.split('/')[-1]
                # Generate a unique filename
                filename = str(uuid.uuid4()) + '.' + ext
                # Decode base64 data
                image_data = base64.b64decode(imgstr)

                # Create a ContentFile from the decoded image data
                image_file = ContentFile(image_data, name=filename)

                # Save the image instance
                image_upload = ImageUpload()
                image_upload.image.save(filename, image_file)

                # Create a unique file name for the new Word document
                timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                doc_filename = f'document.docx'
                doc_path = os.path.join(settings.MEDIA_ROOT, doc_filename)

                x = random.randint(100,1000)
                his_name = f"doc/{doc_filename}"
                his_path = os.path.join(settings.MEDIA_ROOT,his_name)


                # Check if the document exists
                if os.path.exists(doc_path):
                    document = Document(doc_path)
                else:
                    document = Document()

                # Open PIL Image from ContentFile
                pil_image = PILImage.open(image_file)

                # Process image using Gemini API
                response = model.generate_content([pil_image, "act As a vehicle technician, please examine the provided image of the vehicle. Determine if there is any exterior damage.If damage is present, identify and describe the location of the damaged area. Additionally, provide a solution for repairing the damage.If no damage is detected, confirm that the vehicle exterior is in good condition and offer tips for proper maintenance to keep the vehicle in excellent condition.if it is damaged is that should completely replaced or to repair. Give your response within 100 words. do not use bold texts. "])

                # Extract text from the response
                text_response = response.text

                # Add original image to the document
                document.add_picture(image_file,height=Inches(3), width=Inches(6.0))  # Original image

                # Convert PIL image to numpy array
                image_array = np.array(pil_image)

                # Convert to grayscale
                gray_image = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)

                # Apply a temperature colormap
                temperature_image = cv2.applyColorMap(gray_image, cv2.COLORMAP_JET)

                # Save the temperature image to a file
                temperature_image_path = os.path.join(settings.MEDIA_ROOT, f'temperature_image_{timestamp}.png')
                cv2.imwrite(temperature_image_path, temperature_image)

                # Add temperature image to the document
                document.add_picture(temperature_image_path,height=Inches(3), width=Inches(6.0))

                # Add text response to the document
                document.add_paragraph(text_response)

                # Save the document
                document.save(doc_path)
                document.save(his_path)

            except UnrecognizedImageError:
                
                return redirect('upload_image')
            except PermissionError:
                
                return redirect('upload_image')

              # Redirect to success page after processing
        else:
            
            return redirect('upload_image')
    if request.method == 'GET' : 
        a = request.GET.get('lr')
        print(a)
        
    messages.warning(request, 'Exterior images uploaded')
    return render(request,'lab/upload_image.html')

def success(request):
    return render(request,'lab/success.html')


def tires(request):
    if request.method == 'POST':
        lf = request.POST.get('lf')
        rf = request.POST.get('rf')
        lr = request.POST.get('lr')
        rr = request.POST.get('rr')
        model = request.POST.get('model')
        d = {'left front':lf, 'right front':rf, 'left rear':lr, 'right rear':rr}
        data = [
            {
                "Truck number": 7301234,
                "model": "cat 1",
                "tyre pressure (LF, RF, LR, RR)": "(90, 90, 90, 90)",
                "tyre size": "23.5R25",
                "battery make": "ABC",
                "battery replacement date": "2025-09",
                "battery water level": "97%",
                "engine oil color": "Light amber",
                "brake fluid color": "Clear to light yellow"
            },
            {
                "Truck number": 7301235,
                "model": "cat 2",
                "tyre pressure (LF, RF, LR, RR)": "(115, 115, 120, 120)",
                "tyre size": "37.00R57",
                "battery make": "CAT",
                "battery replacement date": "2024-06-12",
                "battery water level": "99%",
                "engine oil color": "Light amber",
                "brake fluid color": "Clear to light yellow"
            },
            {
                "Truck number": 7301236,
                "model": "cat 3",
                "tyre pressure (LF, RF, LR, RR)": "(105, 105, 100, 110)",
                "tyre size": "36.0R56",
                "battery make": "XYZ",
                "battery replacement date": "2024-01-20",
                "battery water level": "98%",
                "engine oil color": "Light amber",
                "brake fluid color": "Clear"
            },
            {
                "Truck number": 7301237,
                "model": "cat 4",
                "tyre pressure (LF, RF, LR, RR)": "(95, 105, 100, 110)",
                "tyre size": "35.0R57",
                "battery make": "DCA",
                "battery replacement date": "2023-12-23",
                "battery water level": "96%",
                "engine oil color": "Light amber",
                "brake fluid color": "Clear"
            },
            {
                "Truck number": 7301238,
                "model": "cat 5",
                "tyre pressure (LF, RF, LR, RR)": "(105, 115, 120, 110)",
                "tyre size": "28.0R57",
                "battery make": "QXC",
                "battery replacement date": "2024-07-09",
                "battery water level": "95%",
                "engine oil color": "Light amber",
                "brake fluid color": "Clear to pale yellow"
            }
        ]
        model = genai.GenerativeModel('gemini-pro')
        prompt = f"""
                    You will be provided with the sample dataset with ideal values of tyre pressure levels and also a new observation dictionary will be provided for a vehicle that had came for service. your task is to analyse the dataset and check whether the pressure levels are correct. if not give instructions whether to reduce or increase the presure in the tyres by addressing potential hazards. Give your response within 100 words. do not use bold texts. 

                    dataset = {data}

                    vehicle data = {d}    
                """
        doc_filename = f'document.docx'
        doc_path = os.path.join(settings.MEDIA_ROOT, doc_filename)

        if os.path.exists(doc_path):
            document = Document(doc_path)
        else:
            document = Document()

        
        response = model.generate_content(prompt)
        result = response.text
        print(result)

        document.add_paragraph(result)
        document.save(doc_path)
        
    return render(request,'lab/tires.html') 



def battery(request):
    if request.method == 'POST':
        # Get the base64-encoded image data from the form
        image_data_url = request.POST.get('image_data')

        if image_data_url:
            try:
                # Decode base64 data
                format, imgstr = image_data_url.split(';base64,')
                ext = format.split('/')[-1]
                # Generate a unique filename
                filename = str(uuid.uuid4()) + '.' + ext
                # Decode base64 data
                image_data = base64.b64decode(imgstr)

                # Create a ContentFile from the decoded image data
                image_file = ContentFile(image_data, name=filename)

                # Save the image instance
                image_upload = ImageUpload()
                image_upload.image.save(filename, image_file)

                # Create a unique file name for the new Word document
                timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                doc_filename = f'document.docx'
                doc_path = os.path.join(settings.MEDIA_ROOT, doc_filename)

                # Check if the document exists
                if os.path.exists(doc_path):
                    document = Document(doc_path)
                else:
                    document = Document()

                # Open PIL Image from ContentFile
                pil_image = PILImage.open(image_file)

                # Process image using Gemini API
                response = model.generate_content([pil_image, "act As a vehicle technician, please examine the provided image of the vehicle's battery. Determine if there is any damage or leakage.If damage or leakage is present, identify and describe the location of the issue. Additionally, provide a solution for repairing or addressing the damage.If no damage or leakage is detected, confirm that the battery is in good condition and offer tips for proper maintenance to keep the battery in excellent condition.if it is damaged is that should completely replaced or to repair.Give your response within 100 words. do not use bold texts. "])

                # Extract text from the response
                text_response = response.text

                # Add original image to the document
                document.add_picture(image_file, width=Inches(5.0))  # Original image

                # Convert PIL image to numpy array
                image_array = np.array(pil_image)

                # Convert to grayscale
                gray_image = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)

                # Apply a temperature colormap
                temperature_image = cv2.applyColorMap(gray_image, cv2.COLORMAP_JET)

                # Save the temperature image to a file
                temperature_image_path = os.path.join(settings.MEDIA_ROOT, f'temperature_image_{timestamp}.png')
                cv2.imwrite(temperature_image_path, temperature_image)

                # Add temperature image to the document
                document.add_picture(temperature_image_path, width=Inches(5.0))

                # Add text response to the document
                document.add_paragraph(text_response)

                # Save the document
                document.save(doc_path)

            except UnrecognizedImageError:
                
                return redirect('battery')
            except PermissionError:
                
                return redirect('battery')

            # Redirect to success page after processing
        else:
            
            return redirect('battery')

        
    messages.warning(request, 'Exterior images uploaded')
    return render(request,'lab/battery.html')

def brake(request):
    if request.method == 'POST':
        # Get the base64-encoded image data from the form
        image_data_url = request.POST.get('image_data')

        if image_data_url:
            try:
                # Decode base64 data
                format, imgstr = image_data_url.split(';base64,')
                ext = format.split('/')[-1]
                # Generate a unique filename
                filename = str(uuid.uuid4()) + '.' + ext
                # Decode base64 data
                image_data = base64.b64decode(imgstr)

                # Create a ContentFile from the decoded image data
                image_file = ContentFile(image_data, name=filename)

                # Save the image instance
                image_upload = ImageUpload()
                image_upload.image.save(filename, image_file)

                # Create a unique file name for the new Word document
                timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                doc_filename = f'document.docx'
                doc_path = os.path.join(settings.MEDIA_ROOT, doc_filename)

                # Check if the document exists
                if os.path.exists(doc_path):
                    document = Document(doc_path)
                else:
                    document = Document()

                # Open PIL Image from ContentFile
                pil_image = PILImage.open(image_file)

                # Process image using Gemini API
                response = model.generate_content([pil_image, "act As a vehicle technician, please examine the color of the vehicle's engine oil.check the condition of the engine oil color.If the engine oil color is not optimal, describe the color and provide the reason for its condition, along with a solution to rectify it. Offer tips for maintaining the engine oil properly to ensure continued optimal performance.Dont tell anything except your opinion on the engine oil codition based on the color of engine oil.if it is damaged is that should completely replaced or to repair.Give your response within 100 words. do not use bold texts. "])

                # Extract text from the response
                text_response = response.text

                # Add original image to the document
                document.add_picture(image_file, width=Inches(5.0))  # Original image

                # Convert PIL image to numpy array
                image_array = np.array(pil_image)

                # Convert to grayscale
                gray_image = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)

                # Apply a temperature colormap
                temperature_image = cv2.applyColorMap(gray_image, cv2.COLORMAP_JET)

                # Save the temperature image to a file
                temperature_image_path = os.path.join(settings.MEDIA_ROOT, f'temperature_image_{timestamp}.png')
                cv2.imwrite(temperature_image_path, temperature_image)

                # Add temperature image to the document
                document.add_picture(temperature_image_path, width=Inches(5.0))

                # Add text response to the document
                document.add_paragraph(text_response)

                # Save the document
                document.save(doc_path)

            except UnrecognizedImageError:
                
                return redirect('brake')
            except PermissionError:
                
                return redirect('brake')

              # Redirect to success page after processing
        else:
            
            return redirect('brake')
        
    messages.warning(request, 'Exterior images uploaded')
    return render(request,'lab/brake.html')

def engine(request):
    if request.method == 'POST':
        # Get the base64-encoded image data from the form
        image_data_url = request.POST.get('image_data')

        if image_data_url:
            try:
                # Decode base64 data
                format, imgstr = image_data_url.split(';base64,')
                ext = format.split('/')[-1]
                # Generate a unique filename
                filename = str(uuid.uuid4()) + '.' + ext
                # Decode base64 data
                image_data = base64.b64decode(imgstr)

                # Create a ContentFile from the decoded image data
                image_file = ContentFile(image_data, name=filename)

                # Save the image instance
                image_upload = ImageUpload()
                image_upload.image.save(filename, image_file)

                # Create a unique file name for the new Word document
                timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
                doc_filename = f'document.docx'
                pdf_filename = f'document.pdf'
                doc_path = os.path.join(settings.MEDIA_ROOT, doc_filename)
        

                # Check if the document exists
                if os.path.exists(doc_path):
                    document = Document(doc_path)
                else:
                    document = Document()

                # Open PIL Image from ContentFile
                pil_image = PILImage.open(image_file)

                # Process image using Gemini API
                response = model.generate_content([pil_image, "act As a vehicle technician, please examine the color of the vehicle's break fluid color.check the condition of the break fluid color.If the break fluid color is not optimal, describe the color and provide the reason for its condition, along with a solution to rectify it. Offer tips for maintaining the break fluid color properly to ensure continued optimal performance.Dont tell anything except your opinion on the break fluid codition based on the color of break fluid.if it is damaged is that should completely replaced or to repair.Give your response within 100 words. do not use bold texts. "])

                # Extract text from the response
                text_response = response.text

                # Add original image to the document
                document.add_picture(image_file, width=Inches(5.0))  # Original image

                # Convert PIL image to numpy array
                image_array = np.array(pil_image)

                # Convert to grayscale
                gray_image = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)

                # Apply a temperature colormap
                temperature_image = cv2.applyColorMap(gray_image, cv2.COLORMAP_JET)

                # Save the temperature image to a file
                temperature_image_path = os.path.join(settings.MEDIA_ROOT, f'temperature_image_{timestamp}.png')
                cv2.imwrite(temperature_image_path, temperature_image)

                # Add temperature image to the document
                document.add_picture(temperature_image_path, width=Inches(5.0))

                # Add text response to the document
                document.add_paragraph(text_response)

                # Save the document
                document.save(doc_path)
                

            except UnrecognizedImageError:
                
                return redirect('engine')
            except PermissionError:
                
                return redirect('engine')

              # Redirect to success page after processing
        else:
            
            return redirect('engine')
        
    messages.warning(request, 'Exterior images uploaded')
    return render(request,'lab/engine.html')

def exterior_ok(request):
    return redirect('tires')

def tyre_ok(request):
    return redirect('brake')

def brake_ok(request):
    return redirect('battery')

def battery_ok(request):
    return redirect('engine')

def qr(request):
    return render(request,'lab/qr.html')

def home(request):
    return render(request,'lab/index.html')

def inspection(request):
    return render(request,'lab/main_dashboard.html')

